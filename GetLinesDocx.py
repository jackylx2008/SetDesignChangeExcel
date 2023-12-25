from docx import Document


class GetLinesDocx:
    def __init__(self, docx_path: str):
        self._proceed_doc = Document(docx_path)
        self._lines = set()

    def read_docx(self):
        for table in self._proceed_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._lines.add(paragraph.text)
                        # for run in paragraph.runs:
                        #     print(run.text)

    def get_lines(self):
        return self._lines

        # for paragraph in self._proceed_doc.paragraphs:
        #     for run in paragraph.runs:
        #         print(run.text)
        # run.text = run.text.replace(old_text, new_text)


def extract_text_from_cell(docx_path, table_index, row_index, column_index):
    doc = Document(docx_path)
    table = doc.tables[table_index]

    cell = table.cell(row_index, column_index)
    paragraphs = cell.paragraphs

    text = "\n".join([paragraph.text for paragraph in paragraphs])
    return text


if __name__ == "__main__":
    docx = r"D:\CloudStation\国会二期\12 北京院-主体\415设计变更\415暖通\展览区\06-03-C2-027-E (拆排烟阀)\word&CAD\06-03-C2-027-E.docx"
    # d = GetLinesDocx(docx)
    # d.read_docx()
    # lines = d.get_lines()
    # for line in lines:
    #     print(line)
    t = extract_text_from_cell(docx, 0, 4, 3)
    print(t)
